"""Word (.docx) generation of the MCAG forms.

Each form can be produced either:
- filled from system data (completed electronically, printable in Word), or
- blank, for printing and manual completion in the field.

Layouts follow the MCAG source documents in source_documents/.
The generated documents always use the institution's exact legal name in
every section and never include borrower publication clauses.
"""
import io

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from mcag.utils import format_cedi, format_date_gh

BLANK = "…………………………………"


def _doc():
    doc = DocxDocument()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    return doc


def _heading(doc, text, level=1, center=False):
    h = doc.add_heading(text, level=level)
    if center:
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return h


def _line(doc, label, value=None, bold_label=True):
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = bold_label
    p.add_run(str(value) if value not in (None, "") else BLANK)
    return p


def _v(value, formatter=None):
    if value in (None, ""):
        return None
    return formatter(value) if formatter else value


def _to_bytes(doc) -> bytes:
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# Loan Application Form
# ---------------------------------------------------------------------------
def loan_application_docx(institution, customer=None, application=None) -> bytes:
    """MCAG Loan Application Form — filled if customer/application given,
    otherwise a blank form for manual completion."""
    c = customer
    a = application
    doc = _doc()
    _heading(doc, institution.legal_name, level=1, center=True)
    _heading(doc, "LOAN APPLICATION FORM", level=2, center=True)

    _heading(doc, "1. PERSONAL DATA", level=3)
    _line(doc, "NAME", c.full_name if c else None)
    _line(doc, "AKA", c.alias if c else None)
    _line(doc, "SEX", c.sex if c else None)
    _line(doc, "DATE OF BIRTH", _v(c.date_of_birth if c else None, format_date_gh))
    _line(doc, "AGE", c.age if c else None)
    _line(doc, "PLACE OF BIRTH", c.place_of_birth if c else None)
    _line(doc, "NATIONALITY", c.nationality if c else None)
    _line(doc, "HOME TOWN & REGION",
          f"{c.home_town or ''} {c.region or ''}".strip() if c else None)
    _line(doc, "GHANA CARD NUMBER", c.ghana_card_number if c else None)
    _line(doc, "ISSUED DATE", _v(c.ghana_card_issue_date if c else None, format_date_gh))
    _line(doc, "EXPIRY DATE", _v(c.ghana_card_expiry_date if c else None, format_date_gh))
    _line(doc, "MARITAL STATUS", c.marital_status if c else None)
    _line(doc, "NO. OF DEPENDANTS", c.dependants if c else None)
    _line(doc, "CYCLE NUMBER", c.cycle_number if c else None)
    _line(doc, "CONTACT NUMBER(S)",
          f"{c.phone_primary or ''} {c.phone_secondary or ''}".strip() if c else None)
    doc.add_paragraph("PLEASE ATTACH A COPY OF GHANA CARD TO FORM").runs[0].italic = True

    _heading(doc, "2. RESIDENTIAL DATA", level=3)
    _line(doc, "HOUSE No", c.house_number if c else None)
    _line(doc, "DIGITAL ADDRESS", c.residential_digital_address if c else None)
    _line(doc, "LOCATION", c.residential_location if c else None)
    _line(doc, "LANDMARK", c.residential_landmark if c else None)
    _line(doc, "DURATION AT CURRENT HOUSE", c.years_at_residence if c else None)
    _line(doc, "ACCOMMODATION TYPE", c.accommodation_status if c else None)
    _line(doc, "NAME OF OWNER (IF RENTED)", c.landlord_name if c else None)
    _line(doc, "RENT EXPIRY DATE", _v(c.rent_expiry_date if c else None, format_date_gh))
    doc.add_paragraph("PLEASE ATTACH A COPY OF RENT AGREEMENT AND/OR UTILITY BILL "
                      "TO FORM").runs[0].italic = True

    _heading(doc, "3. EMPLOYMENT DATA — SELF-EMPLOYED", level=3)
    _line(doc, "NAME OF BUSINESS", c.business_name if c else None)
    _line(doc, "TYPE OF BUSINESS", c.business_type if c else None)
    _line(doc, "YEARS IN BUSINESS", c.years_in_business if c else None)
    _line(doc, "BUSINESS LOCATION", c.business_location if c else None)
    _line(doc, "LANDMARK", c.business_landmark if c else None)
    _line(doc, "YEARS AT CURRENT LOCATION", c.years_at_business_location if c else None)
    _line(doc, "TYPE OF PREMISES", c.premises_type if c else None)
    _line(doc, "PREMISES STATUS", c.premises_status if c else None)
    _line(doc, "ESTIMATED DAILY SALES",
          _v(c.estimated_daily_sales if c else None, format_cedi))
    _line(doc, "ESTIMATED DAILY EXPENSES",
          _v(c.estimated_daily_expenses if c else None, format_cedi))
    _line(doc, "ESTIMATED WORKING CAPITAL",
          _v(c.estimated_working_capital if c else None, format_cedi))
    _line(doc, "NUMBER OF EMPLOYEES", c.number_of_employees if c else None)
    _line(doc, "OTHER INCOME", _v(c.other_income if c else None, format_cedi))
    doc.add_paragraph("PLEASE ATTACH A COPY OF BUSINESS OPERATING PERMIT, GRA TAX "
                      "RECEIPT, SALES RECORDS AND INVOICES TO FORM").runs[0].italic = True

    _heading(doc, "4. EMPLOYMENT DATA — SALARIED WORKER", level=3)
    _line(doc, "NAME OF EMPLOYER", c.employer_name if c else None)
    _line(doc, "LOCATION", c.employer_location if c else None)
    _line(doc, "TYPE OF BUSINESS", c.employer_business_type if c else None)
    _line(doc, "POSITION", c.position if c else None)
    _line(doc, "YEARS IN EMPLOYMENT", c.years_employed if c else None)
    _line(doc, "NET SALARY", _v(c.net_monthly_salary if c else None, format_cedi))
    doc.add_paragraph("PLEASE ATTACH A COPY OF EMPLOYER'S UNDERTAKING LETTER AND "
                      "LAST 3 MONTHS' PAYSLIPS TO FORM").runs[0].italic = True

    _heading(doc, "5. BANK DATA", level=3)
    _line(doc, "NAME OF BANK", c.bank_name if c else None)
    _line(doc, "BRANCH", c.bank_branch if c else None)
    _line(doc, "ACCOUNT NAME", c.bank_account_name if c else None)
    _line(doc, "ACCOUNT NUMBER", c.bank_account_number if c else None)
    _line(doc, "MOBILE MONEY", f"{c.momo_provider or ''} {c.momo_number or ''}".strip()
          if c else None)

    _heading(doc, "6. SPOUSE / NEXT OF KIN DATA", level=3)
    _line(doc, "NAME OF SPOUSE", c.spouse_name if c else None)
    _line(doc, "SPOUSE OCCUPATION", c.spouse_occupation if c else None)
    _line(doc, "SPOUSE CONTACT", c.spouse_phone if c else None)
    _line(doc, "NEXT OF KIN", c.next_of_kin_name if c else None)
    _line(doc, "RELATIONSHIP", c.next_of_kin_relationship if c else None)
    _line(doc, "TEL. NO", c.next_of_kin_phone if c else None)

    _heading(doc, "7. LOAN DATA", level=3)
    _line(doc, "TYPE OF LOAN PRODUCT", a.product.name if a else None)
    _line(doc, "PURPOSE OF LOAN", a.loan_purpose if a else None)
    _line(doc, "AMOUNT REQUESTED", _v(a.amount_requested if a else None, format_cedi))
    _line(doc, "PROPOSED PERIOD",
          f"{a.proposed_tenure} ({a.repayment_frequency})" if a else None)
    _line(doc, "MODE OF PAYMENT", a.proposed_payment_method if a else None)
    _line(doc, "PROPOSED COLLATERAL", a.proposed_collateral if a else None)

    _heading(doc, "8. LOAN APPLICATION AND AUTHORISATION", level=3)
    doc.add_paragraph(
        "I hereby apply for the loan amount stated above in Section (7) of this "
        "application form. The information I have provided hereunder is to the "
        "best of my knowledge true and correct. I authorise "
        f"{institution.legal_name} to verify the correctness of this information "
        "and also obtain additional information it deems necessary in evaluating "
        "my loan application. I accept that my application fee is not refundable "
        "if my application should be declined. I also agree that applying for "
        "this loan is not a guarantee that the loan will be automatically granted.")
    _line(doc, "SIGN / THUMB PRINT OF APPLICANT", None)
    _line(doc, "DATE", _v(a.date_signed if a else None, format_date_gh))
    return _to_bytes(doc)


# ---------------------------------------------------------------------------
# Offer Letter
# ---------------------------------------------------------------------------
def offer_letter_docx(institution, offer=None, application=None, calc=None) -> bytes:
    doc = _doc()
    _heading(doc, institution.legal_name, level=1, center=True)
    doc.add_paragraph(institution.office_address or "")
    doc.add_paragraph()
    customer = application.customer if application else None
    doc.add_paragraph(customer.full_name.upper() if customer else BLANK)
    doc.add_paragraph(customer.residential_location or "" if customer else BLANK)
    p = doc.add_paragraph()
    p.add_run("OBJECT: OFFER OF LOAN FACILITY "
              "(This is neither a Contract nor a commitment to Lend)").bold = True
    doc.add_paragraph(f"Dear {customer.full_name if customer else BLANK},")
    product_name = application.product.name if application else BLANK
    doc.add_paragraph(
        f"Thank you for your recent {product_name} application. After reviewing "
        f"your information and assessment, we are pleased to offer you a "
        f"{product_name} facility based on the following terms and conditions:")

    def money_or_blank(key):
        return format_cedi(calc[key]) if calc else BLANK

    _line(doc, "AMOUNT FINANCED", money_or_blank("gross_amount_financed"))
    _line(doc, "TENURE",
          f"{calc['tenure']} {calc['frequency']} instalments" if calc else BLANK)
    _line(doc, "ANNUAL PERCENTAGE RATE", f"{calc['apr']}%" if calc else BLANK)
    _line(doc, "INTEREST CHARGE", money_or_blank("total_interest"))
    _line(doc, "NUMBER OF PAYMENTS", calc["number_of_instalments"] if calc else BLANK)
    _line(doc, "AMOUNT OF PAYMENTS", money_or_blank("instalment_amount"))
    _line(doc, "WHEN PAYMENTS DUE", calc["frequency"].title() if calc else BLANK)
    _line(doc, "TOTAL OF PAYMENTS", money_or_blank("total_repayment"))
    _line(doc, "PROCESSING FEES", money_or_blank("total_fees"))
    _line(doc, "AMOUNT RECEIVED", money_or_blank("net_amount_received"))

    if calc:
        doc.add_paragraph()
        doc.add_paragraph("REPAYMENT SCHEDULE").runs[0].bold = True
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, h in enumerate(["#", "Due Date", "Principal", "Interest", "Total Due"]):
            hdr[i].text = h
        for row in calc["instalments"]:
            cells = table.add_row().cells
            cells[0].text = str(row["number"])
            cells[1].text = str(row["due_date"])
            cells[2].text = format_cedi(row["principal_due"])
            cells[3].text = format_cedi(row["interest_due"])
            cells[4].text = format_cedi(row["total_due"])

    doc.add_paragraph()
    _line(doc, "COLLATERAL", application.proposed_collateral if application else BLANK)
    doc.add_paragraph(
        "LATE CHARGE: Late payment penalties apply as configured in the loan "
        "product and stated in the loan agreement.")
    doc.add_paragraph(
        "PREPAYMENT: The Borrower has the right to repay or retire this loan "
        "facility before the expiration of its tenor subject to the early "
        "settlement terms of the product.")
    _line(doc, "OFFER VALID UNTIL",
          _v(offer.offer_expiry_date if offer else None, format_date_gh))
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph(institution.manager_name or institution.proprietor_name or "")
    doc.add_paragraph(institution.legal_name)

    _heading(doc, "ACCEPTANCE OF OFFER", level=3)
    doc.add_paragraph(
        f"I am pleased to accept the offer of {product_name} facility of "
        f"{money_or_blank('gross_amount_financed')} given by {institution.legal_name} "
        "and the terms and conditions stated in the Offer Letter. Payment will "
        "be made as outlined above.")
    _line(doc, "BORROWER'S NAME", customer.full_name.upper() if customer else None)
    _line(doc, "Sign", None)
    _line(doc, "Date", None)
    return _to_bytes(doc)


# ---------------------------------------------------------------------------
# Loan Agreement
# ---------------------------------------------------------------------------
def loan_agreement_docx(institution, agreement=None, application=None, calc=None) -> bytes:
    doc = _doc()
    _heading(doc, "LOAN AGREEMENT", level=1, center=True)
    customer = application.customer if application else None
    lender = institution.legal_name  # exact legal name in EVERY section
    agreement_date = format_date_gh(agreement.agreement_date) if agreement else BLANK

    doc.add_paragraph(
        f"This LOAN AGREEMENT is made this {agreement_date} between {lender} "
        f"(hereinafter referred to as “THE LENDER”) whose address is "
        f"{institution.office_address or BLANK} ({institution.digital_address or ''}) AND "
        f"{customer.full_name if customer else BLANK} of "
        f"{(customer.residential_digital_address or customer.residential_location or BLANK) if customer else BLANK}, "
        f"in the {(customer.region or BLANK) if customer else BLANK} Region of the "
        "Republic of Ghana (referred to as THE BORROWER).")

    _heading(doc, "1. PROMISE TO PAY", level=3)
    if calc:
        doc.add_paragraph(
            f"Within {calc['tenure']} {calc['frequency']} instalments, the Borrower "
            f"promises to pay to {lender} the sum of "
            f"{format_cedi(calc['total_repayment'])}. The breakdown is as stated below.")
        _line(doc, "Principal Amount", format_cedi(calc["gross_amount_financed"]))
        _line(doc, "Interest Charge", format_cedi(calc["total_interest"]))
        _line(doc, "Total of Payments", format_cedi(calc["total_repayment"]))
    else:
        doc.add_paragraph(f"The Borrower promises to pay to {lender} the sum of {BLANK}.")
        _line(doc, "Principal Amount", None)
        _line(doc, "Interest Charge", None)
        _line(doc, "Total of Payments", None)

    _heading(doc, "2. TERMS OF REPAYMENT", level=3)
    if calc:
        doc.add_paragraph(
            f"The BORROWER will repay the Loan Principal and Interest in "
            f"{calc['number_of_instalments']} instalments starting on "
            f"{calc['first_due_date']} and ending on {calc['final_due_date']} "
            "according to the attached repayment schedule.")
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, h in enumerate(["#", "Due Date", "Principal", "Interest", "Total Due"]):
            hdr[i].text = h
        for row in calc["instalments"]:
            cells = table.add_row().cells
            cells[0].text = str(row["number"])
            cells[1].text = str(row["due_date"])
            cells[2].text = format_cedi(row["principal_due"])
            cells[3].text = format_cedi(row["interest_due"])
            cells[4].text = format_cedi(row["total_due"])
    else:
        doc.add_paragraph(f"The BORROWER will repay the Loan in {BLANK} instalments "
                          f"starting on {BLANK} and ending on {BLANK}.")

    _heading(doc, "3. EARLY RETIREMENT OF PRINCIPAL", level=3)
    doc.add_paragraph(
        "The Borrower shall be at liberty to retire this loan facility before "
        "the expiration of its tenor subject to the early settlement terms of "
        "the approved loan product.")

    _heading(doc, "4. PENALTY", level=3)
    doc.add_paragraph(
        "In the event the Borrower fails to make payment of an instalment on "
        "the agreed due date, the Borrower shall be liable to pay the penalty "
        "configured for this product and stated in the offer letter. Penalties "
        "apply to the overdue instalment only and are capped by institutional "
        "policy. Interest is not charged on penalties.")

    _heading(doc, "5. DEFAULT", level=3)
    doc.add_paragraph(
        "The Borrower falls into default if the Borrower fails to comply with "
        "the terms of payment. In the event of a default, the total amount "
        "outstanding on this loan facility inclusive of interest, penalty and "
        "other charges for the period agreed herein shall be due immediately "
        f"and {lender} shall be at liberty to demand payment of same.")

    _heading(doc, "6. COLLATERAL", level=3)
    collateral_desc = (application.proposed_collateral if application else None) or (
        "The Borrower hypothecates present and future assets of the business "
        "and household as security for this loan")
    doc.add_paragraph(
        f"{collateral_desc}. Security shall be realised with due regard to the "
        "laws of Ghana to defray any amount outstanding on this loan facility.")

    _heading(doc, "7. RECOVERY", level=3)
    doc.add_paragraph(
        f"By executing this contract, the Borrower and the Guarantor(s) authorise "
        f"{lender}, its workers, agents and assigns to use all lawful means "
        "available, including proceedings in court, to recover any amount "
        "outstanding on this loan.")

    _heading(doc, "8. NOTICES AND GOVERNING LAW", level=3)
    doc.add_paragraph(
        "Any notice required to be given under this agreement or under any law "
        "shall be in writing and shall be delivered at the regular place of "
        "business of the Borrower or the Lender, or at the last known residence "
        "of the Borrower. The governing law shall be the laws of the Republic "
        "of Ghana.")

    _heading(doc, "9. LIABILITY OF GUARANTOR", level=3)
    doc.add_paragraph(
        f"The liability of the Guarantor(s) is personal and unlimited. {lender}, "
        "its workers, agents and assigns can attach the personal property of the "
        "Guarantor(s) and proceed against the estate of the Guarantor(s) to "
        "settle any amount outstanding on this loan. The Lender reserves the "
        "right to proceed against either or all guarantors without first having "
        "to realise any security provided under this agreement.")

    _heading(doc, "10. GUARANTOR'S UNDERTAKING", level=3)
    doc.add_paragraph(
        "By appending their signature(s) to this agreement, the guarantor(s) "
        "undertake to pay the entire amount outstanding on this loan inclusive "
        "of penalties and other charges in the event the Borrower fails or is "
        "unable to comply with this agreement for any reason whatsoever "
        "including but not limited to bankruptcy, death etc.")
    _line(doc, "GUARANTOR NAME", None)
    _line(doc, "OCCUPATION", None)
    _line(doc, "RESIDENCE", None)
    _line(doc, "SIGNATURE", None)
    _line(doc, "CONTACT", None)

    _heading(doc, "11. CREDIT BUREAU", level=3)
    doc.add_paragraph(
        "Please note that the information on this facility will be made "
        "available to a licensed Credit Bureau approved by the Bank of Ghana as "
        "required by the Credit Reporting Act 2007, Act 726. Any default "
        f"without satisfactory arrangement with {lender} will also be reported "
        "to an approved Credit Bureau after formal demand has been made to you "
        "thereof.")

    _heading(doc, "12. DATA PROTECTION", level=3)
    doc.add_paragraph(
        f"{lender} processes the Borrower's personal data in accordance with "
        "the Data Protection Act 2012, Act 843, solely for the assessment, "
        "administration and recovery of this facility and for statutory "
        "reporting.")

    _heading(doc, "13. EXECUTION", level=3)
    doc.add_paragraph(
        "IN WITNESS WHEREOF the parties hereto have set their hands the day and "
        "year first above written.")
    language = (agreement.language_explained if agreement else None) or BLANK
    doc.add_paragraph(
        f"SIGNED by {customer.full_name if customer else BLANK} (after the "
        "content has been read over and explained to him/her in the "
        f"{language} language and he/she seemed to perfectly understand same "
        "before signing or making his/her mark) in the presence of:")
    _line(doc, "WITNESS NAME", agreement.witness_name if agreement else None)
    _line(doc, "WITNESS CONTACT", agreement.witness_phone if agreement else None)
    _line(doc, "BORROWER SIGNATURE / THUMBPRINT", None)
    doc.add_paragraph()
    doc.add_paragraph(f"SIGNED for and on behalf of THE LENDER, {lender}, by:")
    _line(doc, "NAME", institution.manager_name or institution.proprietor_name)
    _line(doc, "SIGNATURE", None)
    return _to_bytes(doc)
